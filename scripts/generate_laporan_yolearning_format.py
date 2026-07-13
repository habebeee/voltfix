#!/usr/bin/env python3
"""Generate Laporan Capstone Voltfix — format YoLearning."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from pathlib import Path


FOOTER_TEXT = "Laporan Capstone Project — Prodi Sistem Informasi, Universitas Esa Unggul"
OUTPUT = Path("/Users/habiburrahman/perkuliahan/Voltfix/Laporan_Capstone_Project_Voltfix.docx")
OUTPUT_DOWNLOADS = Path("/Users/habiburrahman/Downloads/Laporan_Capstone_Project_Voltfix.docx")


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5

    for level, size in [(1, 14), (2, 13), (3, 12)]:
        name = f"Heading {level}"
        h = doc.styles[name]
        h.font.name = "Times New Roman"
        h.font.bold = True
        h.font.size = Pt(size)
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        h.paragraph_format.line_spacing = 1.5


def add_header_footer(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hr = hp.add_run(FOOTER_TEXT)
    hr.font.name = "Times New Roman"
    hr.font.size = Pt(9)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(FOOTER_TEXT)
    fr.font.name = "Times New Roman"
    fr.font.size = Pt(9)

    # page number
    run = fp.add_run("   |   Halaman ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    run._r.append(fld)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    run = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    run._r.append(instr)
    run = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    run._r.append(fld)
    run = p.add_run("(Klik kanan → Perbarui Bidang → Perbarui Seluruh Daftar Isi)")
    run.font.italic = True
    run.font.size = Pt(10)
    run = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    run._r.append(fld)


def p(doc, text, center=False, bold=False, justify=True):
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(4)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(item, style="List Bullet")
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(item, style="List Number")
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    doc.add_paragraph()
    return t


def screenshot_placeholder(doc, gambar_no, judul, petunjuk):
    doc.add_heading(f"Lampiran C.{gambar_no} {judul}", level=3)
    p(doc, f"[SISIPKAN TANGKAPAN LAYAR DI SINI]", center=True, bold=True)
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = box.add_run("┌" + "─" * 50 + "┐\n│" + " " * 18 + "AREA SCREENSHOT" + " " * 17 + "│\n└" + "─" * 50 + "┘")
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(128, 128, 128)
    p(doc, f"Gambar {gambar_no}. {judul}.", center=True)
    p(doc, f"Petunjuk screenshot: {petunjuk}", justify=True)
    doc.add_paragraph()


def build():
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    # === SAMBUNG ===
    p(doc, "LAPORAN", center=True, bold=True)
    p(doc, "CAPSTONE PROJECT", center=True, bold=True)
    doc.add_paragraph()
    p(doc, "JUDUL PROYEK", center=True, bold=True)
    p(doc, "PENGEMBANGAN SISTEM INFORMASI MANAJEMEN SERVIS", center=True, bold=True)
    p(doc, "PERBAIKAN ELEKTRONIK BERBASIS WEB (VOLTFIX)", center=True, bold=True)
    doc.add_paragraph()
    p(doc, "Disusun oleh:", center=True)
    p(doc, "Nama", center=True)
    p(doc, "[Nama Lengkap]   —   NIM: [XXXXXXXXXX]", center=True)
    doc.add_paragraph()
    p(doc, "Program Studi Sistem Informasi", center=True)
    p(doc, "Fakultas Ilmu Komputer  —  Universitas Esa Unggul", center=True)
    doc.add_paragraph()
    p(doc, "Semester / TA", center=True)
    p(doc, "Genap  /  2025–2026", center=True)
    p(doc, "Jalur Capstone", center=True)
    p(doc, "[ X ] Jalur Web       [ ] Jalur Mobile       [ ] Jalur ERP", center=True)
    p(doc, "Dosen Pembimbing", center=True)
    p(doc, "[Nama Dosen Pembimbing, Gelar]", center=True)
    doc.add_page_break()

    # === PENGESAHAN ===
    p(doc, "LEMBAR PENGESAHAN", center=True, bold=True)
    p(doc, "LAPORAN CAPSTONE PROJECT", center=True, bold=True)
    p(doc, "PENGEMBANGAN SISTEM INFORMASI MANAJEMEN SERVIS PERBAIKAN ELEKTRONIK BERBASIS WEB (VOLTFIX)", center=True)
    doc.add_paragraph()
    table(doc, ["", ""], [
        ["Nama", "[Nama Lengkap]"],
        ["Jalur", "[ X ] Web   [ ] Mobile   [ ] ERP"],
        ["Semester / TA", "Genap  /  2025–2026"],
    ])
    p(doc, "Mengetahui,\t\t\t\tMenyetujui,")
    p(doc, "Koordinator Capstone Project\t\tDosen Pembimbing")
    p(doc, "[Nama Koordinator]\t\t\t[Nama Dosen Pembimbing, Gelar]")
    p(doc, "NIDN: ______________\t\t\tNIDN: ______________")
    p(doc, "Tanggal Pengesahan\tJakarta, __________________ 2026")
    p(doc, "Nilai Akhir\t__________ ( ________________________________ )")
    doc.add_page_break()

    # === ABSTRAK ===
    p(doc, "ABSTRAK", center=True, bold=True)
    p(doc, (
        "Voltfix adalah aplikasi web manajemen servis perbaikan perangkat elektronik yang "
        "dikembangkan untuk mendigitalkan siklus tiket servis mulai dari pengajuan pelanggan, "
        "persetujuan admin, penugasan teknisi, pelaksanaan di lapangan, hingga penutupan "
        "setelah penilaian. Permasalahan utama yang melatarbelakangi proyek ini adalah "
        "pengelolaan permintaan servis yang masih bergantung pada pesan singkat, telepon, "
        "dan catatan terpisah sehingga status pekerjaan sulit dilacak dan pelanggan tidak "
        "selalu memperoleh informasi tepat waktu. Sistem dikembangkan menggunakan Laravel 12, "
        "Blade, Tailwind CSS, Filament V3, MariaDB, Docker Compose, dan Nginx. Fitur utama "
        "meliputi autentikasi berbasis peran, portal pelanggan, portal teknisi, panel "
        "administrasi, pengaturan halaman utama, reset kata sandi melalui surel Resend, "
        "serta notifikasi WhatsApp otomatis melalui Fonnte API pada tahapan penting proses "
        "servis. Hasil pengembangan menunjukkan bahwa Voltfix dapat menjadi platform servis "
        "elektronik yang terstruktur dan dapat di-hosting pada VPS. Sistem ini juga "
        "menyediakan ruang pengembangan lanjutan seperti penawaran harga, pembayaran daring, "
        "dan manajemen suku cadang apabila kebutuhan operasional semakin kompleks."
    ))
    p(doc, "Kata Kunci: Voltfix, servis elektronik, manajemen tiket, Laravel, notifikasi WhatsApp.")
    doc.add_paragraph()
    p(doc, "ABSTRACT", center=True, bold=True)
    p(doc, (
        "Voltfix is a web-based electronic repair service management application developed to "
        "digitize the service ticket lifecycle from customer submission, admin approval, "
        "technician assignment, field execution, to closure after rating. The main problem "
        "addressed is fragmented service request handling through messaging apps and manual "
        "records, which makes job status difficult to track. The system was built using "
        "Laravel 12, Blade, Tailwind CSS, Filament V3, MariaDB, Docker Compose, and Nginx. "
        "Key features include role-based authentication, customer and technician portals, "
        "admin panel, homepage settings, email password reset via Resend, and automatic "
        "WhatsApp notifications via Fonnte API at critical service stages. The results "
        "demonstrate that Voltfix is a structured and deployable electronic service platform."
    ))
    p(doc, "Keywords: Voltfix, electronic repair, ticket management, Laravel, WhatsApp notification.")
    doc.add_page_break()

    # === DAFTAR ISI ===
    p(doc, "DAFTAR ISI", center=True, bold=True)
    add_toc(doc)
    doc.add_page_break()

    # === DAFTAR GAMBAR ===
    p(doc, "DAFTAR GAMBAR", center=True, bold=True)
    gambar_list = [
        "Halaman Utama (Landing Page)",
        "Halaman Masuk (Login)",
        "Halaman Daftar (Register)",
        "Halaman Lupa Kata Sandi",
        "Dasbor Pelanggan",
        "Formulir Buat Tiket Servis",
        "Detail Tiket Pelanggan",
        "Formulir Penilaian Teknisi",
        "Dasbor Teknisi",
        "Detail Tiket Teknisi",
        "Panel Admin — Daftar Tiket",
        "Panel Admin — Persetujuan dan Penugasan Teknisi",
        "Panel Admin — Kelola Halaman Utama",
        "Panel Admin — Dasbor Statistik",
        "Contoh Notifikasi WhatsApp ke Pelanggan",
        "Diagram Alur Status Tiket Servis",
        "Diagram Use Case Voltfix",
    ]
    for i, g in enumerate(gambar_list, 1):
        p(doc, f"{i}. Gambar {i}. {g}")
    doc.add_page_break()

    # === DAFTAR TABEL ===
    p(doc, "DAFTAR TABEL", center=True, bold=True)
    tabel_list = [
        "Identitas Dokumen",
        "Definisi Istilah",
        "Aktor Sistem",
        "Ruang Lingkup dan Batasan",
        "Teknologi yang Digunakan",
        "Modul Sistem",
        "Entitas Data Utama",
        "Pemicu Notifikasi WhatsApp",
        "Skenario Pengujian",
        "Risiko dan Mitigasi",
    ]
    for i, t in enumerate(tabel_list, 1):
        p(doc, f"{i}. Tabel {i}. {t}")
    doc.add_page_break()

    # ===== BAB 1 =====
    doc.add_heading("BAB 1 PENDAHULUAN", level=1)
    doc.add_heading("1.1 Latar Belakang", level=2)
    p(doc, (
        "Perkembangan layanan digital mendorong usaha jasa perbaikan elektronik untuk "
        "menyediakan proses pengajuan dan pelacakan servis yang lebih terstruktur. "
        "Televisi, telepon genggam, dan komputer portabel merupakan perangkat yang "
        "sering memerlukan perbaikan, namun banyak penyedia layanan masih mencatat "
        "permintaan melalui pesan singkat, panggilan telepon, atau lembar kerja manual."
    ))
    p(doc, (
        "Kondisi tersebut menimbulkan antrian pekerjaan yang sulit diatur, penugasan "
        "teknisi yang tidak selalu sesuai keahlian kategori perangkat, serta "
        "keterbatasan transparansi status bagi pelanggan. Riwayat servis dan penilaian "
        "kinerja teknisi pun belum terdokumentasi secara sistematis."
    ))
    p(doc, (
        "Voltfix dikembangkan sebagai platform manajemen servis perbaikan elektronik "
        "berbasis web yang menghubungkan pelanggan, teknisi lapangan, admin, dan manajer "
        "dalam satu alur digital. Pelanggan terdaftar dapat mengajukan tiket servis untuk "
        "kategori TV, HP, atau laptop; admin memproses antrian dan menugaskan teknisi; "
        "teknisi memperbarui progres pekerjaan; pelanggan memberikan penilaian setelah "
        "servis selesai."
    ))
    p(doc, (
        "Sistem turut mengintegrasikan reset kata sandi melalui surel dan notifikasi "
        "WhatsApp otomatis agar pelanggan menerima konfirmasi pada tahapan penting "
        "tanpa harus membuka aplikasi terus-menerus. Pengembangan mengacu pada Business "
        "Requirements Document dan Product Requirements Document Voltfix versi 1.2."
    ))

    doc.add_heading("1.2 Rumusan Masalah", level=2)
    numbered(doc, [
        "Bagaimana merancang aplikasi web manajemen servis elektronik yang memiliki alur tiket terstruktur dari pengajuan hingga penutupan?",
        "Bagaimana menyediakan portal berbeda bagi pelanggan, teknisi, admin, dan manajer sesuai kewenangan masing-masing?",
        "Bagaimana membangun mekanisme pelacakan status dan notifikasi WhatsApp yang menjaga konsistensi informasi bagi pelanggan?",
        "Bagaimana mengintegrasikan reset kata sandi melalui surel dan notifikasi otomatis ke perangkat genggam pelanggan?",
        "Bagaimana memverifikasi fungsi kritis sistem dan menyiapkan penerapannya pada lingkungan produksi (VPS)?",
    ])

    doc.add_heading("1.3 Tujuan Penelitian", level=2)
    numbered(doc, [
        "Menghasilkan aplikasi web Voltfix sebagai platform manajemen servis perbaikan elektronik berbasis Laravel.",
        "Menyediakan alur tiket terstruktur berdasarkan kategori perangkat, status pekerjaan, dan penugasan teknisi.",
        "Menyediakan portal pelanggan, portal teknisi, dan panel administrasi berbasis peran pengguna.",
        "Menerapkan notifikasi WhatsApp otomatis melalui Fonnte API dan reset kata sandi melalui surel Resend.",
        "Menyediakan pengaturan halaman utama, jejak audit perubahan status, dan dasbor operasional admin.",
        "Mendokumentasikan implementasi, pengujian, dan penerapan pada server VPS.",
    ])

    doc.add_heading("1.4 Manfaat Penelitian", level=2)
    p(doc, "1.4.1 Bagi pelanggan, Voltfix memudahkan pengajuan servis, pemantauan status, penerimaan notifikasi WhatsApp, dan pemberian penilaian teknisi melalui satu akun.")
    p(doc, "1.4.2 Bagi teknisi, Voltfix menyediakan daftar pekerjaan terstruktur dan formulir pembaruan status sesuai tahapan operasional di lapangan.")
    p(doc, "1.4.3 Bagi admin dan manajer, Voltfix menyediakan antrian tiket, penugasan berbasis keahlian, pengelolaan pengguna, serta dasbor statistik operasional.")
    p(doc, "1.4.4 Bagi pengembang dan institusi, proyek ini menjadi implementasi nyata rekayasa perangkat lunak web, basis data relasional, otorisasi berbasis peran, integrasi layanan eksternal, dan deployment VPS.")

    doc.add_heading("1.5 Ruang Lingkup dan Batasan", level=2)
    p(doc, "Ruang lingkup proyek mencakup website pelanggan, portal teknisi, panel admin, notifikasi surel dan WhatsApp, serta deployment web. Batasan dibuat agar pengembangan realistis dan selaras BRD/PRD.")
    table(doc, ["Kategori", "Termasuk", "Tidak Termasuk"], [
        ["Servis", "Tiket TV/HP/Laptop, foto kerusakan, alamat, jadwal preferensi, rating", "POS, inventori retail, katalog suku cadang"],
        ["Notifikasi", "Surel reset kata sandi (Resend), WhatsApp otomatis (Fonnte)", "Aplikasi seluler native, notifikasi real-time in-app"],
        ["Admin", "Tiket, teknisi, pengguna, peran, log aktivitas, halaman utama", "Sistem ERP atau marketplace eksternal"],
        ["Pembayaran", "Nomor faktur sebagai ID pelacakan", "Gateway pembayaran daring, quotation formal"],
        ["Integrasi", "Fonnte API, Resend, Docker, VPS", "API publik REST, integrasi IoT perangkat"],
    ])
    p(doc, "Tabel 1.5. Ruang lingkup dan batasan — Sumber: BRD Voltfix v1.2, 2026")
    doc.add_page_break()

    # ===== BAB 2 =====
    doc.add_heading("BAB 2 LANDASAN TEORI DAN KAJIAN KEBUTUHAN", level=1)
    doc.add_heading("2.1 Sistem Informasi Manajemen Servis", level=2)
    p(doc, "Sistem informasi manajemen servis mengintegrasikan data permintaan, pelaksanaan, dan umpan balik dalam satu platform. Nilainya terletak pada keterhubungan antardata: profil pelanggan, tiket, teknisi, log status, penilaian, dan log notifikasi saling merujuk sehingga setiap perubahan dapat ditelusuri.")

    doc.add_heading("2.2 Aplikasi Web dan Arsitektur Berlapis", level=2)
    p(doc, "Aplikasi web dapat diakses melalui peramban tanpa pemasangan khusus. Arsitektur berlapis memisahkan antarmuka, logika aplikasi, akses data, dan infrastruktur. Voltfix menggunakan pola Model–View–Controller pada Laravel dengan Blade dan Tailwind CSS untuk antarmuka publik.")

    doc.add_heading("2.3 Manajemen Tiket dan Alur Status", level=2)
    p(doc, "Manajemen tiket menggunakan status eksplisit untuk menggambarkan progres servis. Setiap transisi status dicatat pada log audit sehingga admin, teknisi, dan pelanggan memiliki referensi yang sama. Pendekatan ini mengurangi miskomunikasi dibandingkan pencatatan informal.")

    doc.add_heading("2.4 Notifikasi Otomatis dan Layanan Surel", level=2)
    p(doc, "Notifikasi otomatis mempercepat penyebaran informasi kepada pelanggan pada momen penting. Voltfix menggunakan Fonnte API untuk WhatsApp dan Resend untuk surel reset kata sandi. Kredensial layanan disimpan pada variabel lingkungan agar aman.")

    doc.add_heading("2.5 Panel Administrasi Berbasis Filament", level=2)
    p(doc, "Panel administrasi memungkinkan pengelola mengatur data operasional tanpa mengubah basis data secara manual. Filament V3 menyediakan resource, formulir, tabel, aksi, dan widget yang terintegrasi sehingga proses approve, reject, dan assign teknisi dapat dilakukan efisien.")

    doc.add_heading("2.6 Teknologi Web yang Digunakan", level=2)
    table(doc, ["Teknologi", "Peran dalam Sistem"], [
        ["Laravel 12", "Kerangka kerja backend: routing, controller, model, migrasi, layanan, autentikasi"],
        ["PHP 8.3", "Bahasa pemrograman backend"],
        ["Blade + Tailwind CSS", "Antarmuka landing page, portal pelanggan, portal teknisi"],
        ["Filament V3", "Panel admin: tiket, teknisi, pengguna, pengaturan halaman utama"],
        ["MariaDB", "Basis data relasional"],
        ["Docker Compose", "Lingkungan pengembangan: PHP-FPM, Nginx, MariaDB"],
        ["Nginx", "Server web"],
        ["Resend", "Pengiriman surel reset kata sandi"],
        ["Fonnte API", "Pengiriman notifikasi WhatsApp otomatis"],
        ["Filament Shield", "Pengelolaan peran dan izin (Spatie Permission)"],
        ["Git dan GitHub", "Version control dan sinkronisasi ke VPS"],
    ])
    p(doc, "Tabel 2.6. Teknologi yang digunakan — Sumber: BRD Voltfix, 2026")
    doc.add_page_break()

    # ===== BAB 3 =====
    doc.add_heading("BAB 3 ANALISIS KEBUTUHAN DAN PERANCANGAN SISTEM", level=1)
    doc.add_heading("3.1 Gambaran Umum Sistem", level=2)
    p(doc, "Voltfix memiliki dua sisi utama: sisi pengguna (pengunjung, pelanggan, teknisi) dan sisi pengelola (admin, manajer). Sisi pengguna untuk mengajukan tiket, memantau status, menerima notifikasi, dan memberi penilaian. Sisi pengelola untuk memproses antrian, menugaskan teknisi, mengelola pengguna, dan memantau operasional.")

    doc.add_heading("3.2 Aktor Sistem", level=2)
    table(doc, ["Aktor", "Deskripsi", "Hak Akses Utama"], [
        ["Pengunjung", "Pengguna belum masuk", "Melihat halaman utama, daftar, masuk"],
        ["Pelanggan", "Pengguna terdaftar role CUSTOMER", "Mengajukan tiket, memantau status, memberi penilaian"],
        ["Teknisi", "Petugas lapangan", "Melihat pekerjaan ditugaskan, memperbarui status"],
        ["Manajer", "Pengawas operasional", "Memantau tiket dan teknisi (akses baca dominan)"],
        ["Admin", "Pengelola sistem", "Approve/reject/assign tiket, kelola teknisi, pengguna, halaman utama"],
    ])
    p(doc, "Tabel 3.2. Aktor sistem — Sumber: BRD Voltfix, 2026")

    doc.add_heading("3.3 Kebutuhan Fungsional", level=2)
    table(doc, ["Modul", "Kebutuhan"], [
        ["Halaman Utama", "Menampilkan identitas brand, ringkasan layanan, alur servis, tombol daftar/masuk"],
        ["Autentikasi", "Masuk, daftar, keluar, lupa kata sandi, reset kata sandi, pengalihan portal per peran"],
        ["Portal Pelanggan", "Dasbor statistik, daftar tiket, buat tiket, detail tiket, penilaian teknisi"],
        ["Portal Teknisi", "Dasbor pekerjaan, detail tiket, pembaruan status, edit profil keahlian"],
        ["Panel Admin", "Kelola tiket (approve, reject, assign), teknisi, pengguna, peran, log aktivitas"],
        ["Pengaturan Halaman Utama", "Logo, gambar hero, langkah layanan, gambar kategori TV/HP/Laptop"],
        ["Notifikasi Surel", "Reset kata sandi via Resend dengan notifikasi berbahasa Indonesia"],
        ["Notifikasi WhatsApp", "Pesan otomatis via Fonnte pada pengajuan, approve, reject, assign, selesai"],
        ["Unggah Gambar", "Foto kerusakan tiket (min 1, max 5, max 2 MB) dan gambar halaman utama"],
    ])

    doc.add_heading("3.4 Kebutuhan Nonfungsional", level=2)
    p(doc, "3.4.1 Kegunaan: antarmuka berbahasa Indonesia, alur formulir jelas, dan pesan validasi informatif.")
    p(doc, "3.4.2 Responsivitas: halaman dapat digunakan pada peramban desktop dan genggam.")
    p(doc, "3.4.3 Keamanan: autentikasi sesi, perlindungan CSRF, otorisasi peran, pemeriksaan kepemilikan tiket.")
    p(doc, "3.4.4 Integritas Data: setiap perubahan status tercatat pada ticket_logs; log WhatsApp pada wa_logs.")
    p(doc, "3.4.5 Pemeliharaan: struktur Laravel dengan controller, model, migrasi, layanan, dan Blade.")
    p(doc, "3.4.6 Keandalan: penanganan token Fonnte kosong (dev_logged), timeout API 10 detik, normalisasi nomor telepon.")
    p(doc, "3.4.7 Penerapan: aplikasi berjalan melalui Docker dan dapat dipublikasikan pada VPS.")

    doc.add_heading("3.5 Perancangan Data", level=2)
    table(doc, ["Entitas", "Fungsi"], [
        ["users", "Akun pelanggan, teknisi, admin, manajer; nama, surel, telepon, peran"],
        ["technicians", "Profil teknisi: kategori keahlian, rating rata-rata, ketersediaan, pengalaman"],
        ["tickets", "Tiket servis: nomor faktur, kategori, foto, alamat, jadwal, status"],
        ["ticket_logs", "Jejak audit perubahan status tiket"],
        ["ratings", "Penilaian 1–5 dan ulasan pelanggan terhadap teknisi"],
        ["wa_logs", "Log pengiriman notifikasi WhatsApp: nomor, pesan, status"],
        ["site_settings", "Pengaturan konten halaman utama (pasangan kunci–nilai)"],
    ])
    p(doc, "Tabel 3.5. Entitas data utama — Sumber: BRD Voltfix, 2026")

    doc.add_heading("3.6 Perancangan Alur Status Tiket", level=2)
    p(doc, "Alur status: PENDING → (admin setujui) WAITING_ASSIGNMENT → (admin assign) ASSIGNED → ON_THE_WAY → DIAGNOSIS → REPAIR / WAITING_PART → COMPLETED → (pelanggan rating) CLOSED. Admin dapat menolak tiket PENDING menjadi REJECTED.")
    p(doc, "[SISIPKAN GAMBAR 16 — Diagram Alur Status Tiket Servis di sini]")
    p(doc, "Gambar 16. Diagram alur status tiket servis Voltfix.", center=True)

    doc.add_heading("3.7 Perancangan Notifikasi WhatsApp", level=2)
    table(doc, ["No", "Pemicu", "Method", "Isi Pesan Singkat"], [
        ["1", "Pelanggan mengajukan tiket", "sendTicketSubmitted()", "Konfirmasi tiket diterima + nomor faktur"],
        ["2", "Admin menyetujui tiket", "sendTicketApproved()", "Tiket disetujui, menunggu penugasan"],
        ["3", "Admin menolak tiket", "sendTicketRejected()", "Tiket ditolak + alasan"],
        ["4", "Admin menugaskan teknisi", "sendTicketAssigned()", "Nama teknisi yang ditugaskan"],
        ["5", "Teknisi menyelesaikan servis", "sendTicketCompleted()", "Servis selesai + ajakan penilaian"],
    ])
    p(doc, "Tabel 3.7. Pemicu notifikasi WhatsApp — Sumber: implementasi WhatsAppService, 2026")
    doc.add_page_break()

    # ===== BAB 4 =====
    doc.add_heading("BAB 4 IMPLEMENTASI SISTEM", level=1)
    doc.add_heading("4.1 Implementasi Halaman Utama dan Autentikasi", level=2)
    p(doc, "Halaman utama dibangun dengan Blade dan Tailwind CSS, menampilkan identitas brand, ringkasan layanan, dan alur tiga langkah servis. Konten dinamis diambil dari tabel site_settings sehingga admin dapat memperbarui logo, gambar hero, dan gambar kategori tanpa mengubah kode.")
    p(doc, "Autentikasi menyediakan masuk, daftar, keluar, lupa kata sandi, dan reset kata sandi. Registrasi publik hanya menghasilkan akun CUSTOMER. Setelah masuk, pengguna diarahkan otomatis ke portal sesuai peran. Reset kata sandi menggunakan ResetPasswordNotification berbahasa Indonesia via Resend.")

    doc.add_heading("4.2 Implementasi Portal Pelanggan", level=2)
    p(doc, "Portal pelanggan meliputi dasbor statistik (total, aktif, selesai, menunggu), daftar tiket dengan paginasi, formulir buat tiket, detail tiket, dan formulir penilaian. Formulir buat tiket memvalidasi kategori TV/HP/Laptop, penyebab kerusakan, minimal satu foto, alamat lengkap, dan jadwal preferensi. Nomor faktur otomatis dihasilkan dengan format INV-YYYYMMDD-####.")

    doc.add_heading("4.3 Implementasi Portal Teknisi", level=2)
    p(doc, "Teknisi hanya melihat tiket yang ditugaskan kepadanya. Pembaruan status mengikuti mesin status: ASSIGNED → ON_THE_WAY → DIAGNOSIS → REPAIR atau WAITING_PART → COMPLETED. Setiap pembaruan dicatat pada ticket_logs. Saat status COMPLETED, sistem memanggil WhatsAppService untuk mengirim notifikasi penyelesaian.")

    doc.add_heading("4.4 Implementasi Panel Admin", level=2)
    p(doc, "Panel admin menggunakan Filament V3 dengan TicketResource, TechnicianResource, UserResource, ManageHomePage, widget statistik, Filament Shield, dan activity log. Admin dapat approve, reject, dan assign teknisi langsung dari tabel tiket. Setiap aksi admin memicu notifikasi WhatsApp sesuai Tabel 3.7.")

    doc.add_heading("4.5 Implementasi Layanan WhatsApp (Fonnte API)", level=2)
    p(doc, "WhatsAppService mengirim pesan melalui endpoint Fonnte dengan header Authorization berisi token. Nomor telepon dinormalisasi ke format 62xxxxxxxxxx. Setiap percobaan pengiriman dicatat pada wa_logs dengan status sent, failed, dev_logged, atau skipped. Konfigurasi FONNTE_URL dan FONNTE_TOKEN disimpan pada berkas .env.")

    doc.add_heading("4.6 Implementasi Penerapan (Deployment)", level=2)
    p(doc, "Aplikasi di-deploy menggunakan Docker Compose (PHP-FPM, Nginx, MariaDB) pada VPS. Kode disinkronkan dari repositori GitHub ke direktori ~/apps/voltfix/src. Setelah konfigurasi token Fonnte, perintah php artisan config:clear dijalankan pada container PHP produksi.")
    doc.add_page_break()

    # ===== BAB 5 =====
    doc.add_heading("BAB 5 PENGUJIAN DAN EVALUASI", level=1)
    doc.add_heading("5.1 Metode Pengujian", level=2)
    p(doc, "Pengujian menggunakan pendekatan black-box testing: fungsi diuji berdasarkan kebutuhan BRD/PRD tanpa memeriksa detail kode internal. Fokus pada alur pengguna, validasi input, penyimpanan data, hak akses, pengiriman notifikasi, dan hasil akhir fitur.")

    doc.add_heading("5.2 Skenario Pengujian", level=2)
    table(doc, ["No", "Fitur", "Skenario", "Hasil yang Diharapkan"], [
        ["1", "Halaman Utama", "Pengunjung membuka halaman utama", "Brand, layanan, alur servis, tombol daftar/masuk tampil"],
        ["2", "Daftar", "Pengguna mengisi formulir daftar", "Akun CUSTOMER terbentuk dengan nomor WhatsApp"],
        ["3", "Masuk", "Pengguna masuk dengan surel dan kata sandi valid", "Diarahkan ke portal sesuai peran"],
        ["4", "Buat Tiket", "Pelanggan upload foto dan kirim formulir", "Tiket PENDING, notifikasi WhatsApp terkirim"],
        ["5", "Approve Admin", "Admin menyetujui tiket PENDING", "Status WAITING_ASSIGNMENT, WhatsApp terkirim"],
        ["6", "Assign Teknisi", "Admin menugaskan teknisi sesuai kategori", "Status ASSIGNED, WhatsApp terkirim"],
        ["7", "Update Teknisi", "Teknisi ubah status hingga COMPLETED", "Log tercatat, WhatsApp selesai terkirim"],
        ["8", "Penilaian", "Pelanggan beri rating 1–5", "Rating tersimpan, tiket CLOSED, rating teknisi terbarui"],
        ["9", "Reset Kata Sandi", "Pengguna minta tautan reset via surel", "Surel terkirim via Resend"],
        ["10", "Akses Tidak Sah", "Pelanggan A buka tiket milik pelanggan B", "Respons 403 Forbidden"],
        ["11", "Reject Admin", "Admin tolak tiket dengan alasan", "Status REJECTED, WhatsApp berisi alasan"],
        ["12", "Log WhatsApp", "Cek tabel wa_logs setelah pengajuan tiket", "Status sent dan nomor tercatat"],
    ])
    p(doc, "Tabel 5.2. Skenario pengujian — Sumber: hasil pengujian manual, Juli 2026")

    doc.add_heading("5.3 Evaluasi Hasil", level=2)
    p(doc, "Berdasarkan implementasi dan pengujian, Voltfix telah memenuhi kebutuhan utama BRD dan PRD. Siklus tiket end-to-end berjalan, notifikasi WhatsApp aktif setelah device Fonnte terhubung dan token dikonfigurasi, serta reset kata sandi surel berfungsi. Fitur di luar ruang lingkup—pembayaran daring, quotation, manajemen suku cadang—dicatat sebagai pengembangan lanjutan.")

    doc.add_heading("5.4 Risiko dan Mitigasi", level=2)
    table(doc, ["Risiko", "Dampak", "Mitigasi"], [
        ["Device Fonnte terputus", "Notifikasi WhatsApp gagal", "Pantau status connect; komunikasi manual sementara"],
        ["Kuota Fonnte habis", "Pesan tidak terkirim", "Pantau wa_logs; upgrade paket Fonnte"],
        ["Nomor WhatsApp tidak valid", "Pesan gagal ke pelanggan", "Validasi format 62xxx saat registrasi"],
        ["Surel reset gagal", "Pengguna tidak dapat reset kata sandi", "Verifikasi domain Resend; pesan error jelas"],
        ["Salah assign teknisi", "Pekerjaan tidak sesuai keahlian", "Filter teknisi by skill_category + is_available"],
        ["Kehilangan data", "Riwayat servis hilang", "Backup MariaDB berkala pada VPS"],
    ])
    doc.add_page_break()

    # ===== BAB 6 =====
    doc.add_heading("BAB 6 KESIMPULAN DAN SARAN", level=1)
    doc.add_heading("6.1 Kesimpulan", level=2)
    numbered(doc, [
        "Voltfix berhasil dikembangkan sebagai aplikasi web manajemen servis perbaikan elektronik yang menghubungkan pelanggan, teknisi, admin, dan manajer dalam satu platform.",
        "Sistem telah menyediakan alur utama mulai dari halaman utama, registrasi, pengajuan tiket, persetujuan admin, penugasan teknisi, pembaruan progres, penilaian, hingga penutupan tiket.",
        "Panel admin Filament berhasil mendukung pengelolaan tiket, teknisi, pengguna, peran, log aktivitas, dan pengaturan halaman utama.",
        "Notifikasi WhatsApp otomatis melalui Fonnte API dan reset kata sandi melalui Resend telah aktif, sehingga pelanggan menerima informasi pada tahapan penting proses servis.",
        "Aplikasi dapat di-hosting pada VPS menggunakan Docker Compose, membuktikan sistem siap digunakan sebagai layanan web publik.",
    ])

    doc.add_heading("6.2 Saran", level=2)
    numbered(doc, [
        "Melaksanakan uji penerimaan pengguna formal dengan responden pelanggan, teknisi, dan admin operasional.",
        "Menambah pengujian otomatis (PHPUnit/Pest) pada modul tiket dan layanan WhatsApp.",
        "Mengembangkan modul penawaran harga dan pembayaran daring pada fase berikutnya.",
        "Memperluas notifikasi WhatsApp untuk pembaruan status teknisi di lapangan jika diperlukan operasional.",
        "Menetapkan prosedur pencadangan basis data dan pemantauan kuota Fonnte secara berkala.",
    ])
    doc.add_page_break()

    # ===== DAFTAR PUSTAKA =====
    doc.add_heading("DAFTAR PUSTAKA", level=1)
    refs = [
        "[1] Filament. (2026). Filament Documentation. https://filamentphp.com/docs",
        "[2] Fonnte. (2026). Dokumentasi API Pengiriman Pesan WhatsApp. https://docs.fonnte.com",
        "[3] Laravel. (2026). Laravel 12 Documentation. https://laravel.com/docs",
        "[4] Resend. (2026). Dokumentasi API Surel. https://resend.com/docs",
        "[5] Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill.",
        "[6] Sommerville, I. (2016). Software Engineering (10th ed.). Pearson.",
        "[7] Tim Pengembang Voltfix. (2026). Business Requirements Document Voltfix v1.2. Universitas Esa Unggul.",
        "[8] Tim Pengembang Voltfix. (2026). Product Requirements Document Voltfix. Universitas Esa Unggul.",
    ]
    for ref in refs:
        p(doc, ref, justify=False)

    doc.add_page_break()

    # ===== LAMPIRAN A =====
    doc.add_heading("LAMPIRAN A KEPUTUSAN RUANG LINGKUP", level=1)
    table(doc, ["No", "Keputusan", "Alasan"], [
        ["1", "Pembayaran daring tidak disertakan", "Fokus proyek pada manajemen tiket; gateway memerlukan akun resmi dan pengujian transaksi"],
        ["2", "Quotation/invoice biaya tidak disertakan", "Nomor faktur hanya sebagai ID pelacakan, bukan dokumen keuangan"],
        ["3", "Manajemen suku cadang tidak disertakan", "Status WAITING_PART ada tanpa modul parts terpisah"],
        ["4", "Notifikasi WhatsApp via Fonnte dipilih", "Pelanggan sudah terbiasa dengan WhatsApp; Fonnte menyediakan API sederhana"],
        ["5", "Website di-hosting pada VPS", "Membuktikan sistem dapat berjalan sebagai layanan web publik"],
    ])
    doc.add_page_break()

    # ===== LAMPIRAN B — MATRIKS =====
    doc.add_heading("LAMPIRAN B MATRIKS KETERLACAKAN KEBUTUHAN", level=1)
    table(doc, ["Kode", "Kebutuhan BRD/PRD", "Modul", "Status"], [
        ["F-01", "Autentikasi dan registrasi", "Auth controllers", "Selesai"],
        ["F-02", "Pengajuan tiket servis", "Customer/TicketController", "Selesai"],
        ["F-03", "Approve/reject/assign admin", "TicketResource Filament", "Selesai"],
        ["F-04", "Update status teknisi", "Technician/TicketController", "Selesai"],
        ["F-05", "Penilaian teknisi", "Customer/TicketController@rate", "Selesai"],
        ["F-06", "Reset kata sandi surel", "ForgotPasswordController + Resend", "Aktif"],
        ["F-07", "Notifikasi WhatsApp", "WhatsAppService + Fonnte", "Aktif"],
        ["F-08", "Pengaturan halaman utama", "ManageHomePage", "Selesai"],
        ["F-09", "Dasbor statistik admin", "Filament Widgets", "Selesai"],
    ])
    doc.add_page_break()

    # ===== LAMPIRAN C — SCREENSHOTS =====
    doc.add_heading("LAMPIRAN C TANGKAPAN LAYAR APLIKASI", level=1)
    p(doc, (
        "Lampiran berikut digunakan sebagai dokumentasi antarmuka aplikasi Voltfix yang "
        "telah diimplementasikan. Setiap ruang kosong ditandai agar penulis dapat "
        "menyisipkan tangkapan layar dari website yang sudah berjalan (lokal atau VPS)."
    ))

    screenshots = [
        (1, "Halaman Utama (Landing Page)", "Buka halaman / (welcome) sebagai pengunjung. Tangkap seluruh halaman: logo, hero, layanan TV/HP/Laptop, alur 3 langkah."),
        (2, "Halaman Masuk (Login)", "Buka /login. Tangkap formulir surel, kata sandi, dan tautan daftar/lupa kata sandi."),
        (3, "Halaman Daftar (Register)", "Buka /register. Tangkap formulir nama, surel, nomor WhatsApp (628xxx), kata sandi."),
        (4, "Halaman Lupa Kata Sandi", "Buka /forgot-password. Tangkap formulir permintaan reset surel."),
        (5, "Dasbor Pelanggan", "Masuk sebagai pelanggan. Tangkap /customer/dashboard: statistik tiket (total, aktif, selesai, menunggu)."),
        (6, "Formulir Buat Tiket Servis", "Buka /customer/tickets/create. Tangkap pilihan kategori, upload foto, alamat, jadwal."),
        (7, "Detail Tiket Pelanggan", "Buka detail tiket setelah dibuat. Tangkap nomor faktur, status, riwayat log, foto kerusakan."),
        (8, "Formulir Penilaian Teknisi", "Buka tiket berstatus COMPLETED. Tangkap formulir rating 1–5 dan ulasan."),
        (9, "Dasbor Teknisi", "Masuk sebagai teknisi. Tangkap /technician/dashboard: statistik dan daftar pekerjaan aktif."),
        (10, "Detail Tiket Teknisi", "Buka detail tiket dari sisi teknisi. Tangkap data pelanggan, alamat, tombol update status."),
        (11, "Panel Admin — Daftar Tiket", "Masuk /admin, buka TicketResource. Tangkap tabel tiket dengan badge status dan tombol aksi."),
        (12, "Panel Admin — Persetujuan dan Penugasan", "Tangkap modal/dialog Setujui, Tolak, atau Tugaskan Teknisi."),
        (13, "Panel Admin — Kelola Halaman Utama", "Buka menu Kelola Halaman Utama. Tangkap formulir logo, hero, gambar layanan."),
        (14, "Panel Admin — Dasbor Statistik", "Tangkap widget statistik tiket, pending, chart, rating teknisi."),
        (15, "Contoh Notifikasi WhatsApp", "Tangkap layar HP pelanggan setelah buat tiket/admin approve: pesan dari device Fonnte Voltfix."),
    ]
    for no, judul, petunjuk in screenshots:
        screenshot_placeholder(doc, no, judul, petunjuk)

    doc.add_heading("Lampiran C.16 Diagram Alur Status Tiket", level=3)
    p(doc, "[SISIPKAN DIAGRAM ALUR STATUS TIKET — bisa dari draw.io atau tools UML]", center=True, bold=True)
    p(doc, "Gambar 16. Diagram alur status tiket servis Voltfix.", center=True)
    p(doc, "Petunjuk: Buat diagram alur PENDING → WAITING_ASSIGNMENT → ASSIGNED → ... → COMPLETED → CLOSED, serta cabang REJECTED.", justify=True)
    doc.add_paragraph()

    doc.add_heading("Lampiran C.17 Diagram Use Case", level=3)
    p(doc, "[SISIPKAN DIAGRAM USE CASE — file use_case_voltfix_horizontal.drawio]", center=True, bold=True)
    p(doc, "Gambar 17. Diagram use case Voltfix.", center=True)
    p(doc, "Petunjuk: Ekspor dari draw.io sebagai PNG, sisipkan aktor Pengunjung, Pelanggan, Teknisi, Admin, Manajer.", justify=True)

    doc.save(OUTPUT)
    try:
        import shutil
        shutil.copy2(OUTPUT, OUTPUT_DOWNLOADS)
        print(f"Salin ke: {OUTPUT_DOWNLOADS}")
    except OSError as e:
        print(f"Catatan: salin manual ke Downloads jika perlu ({e})")
    print(f"Berhasil: {OUTPUT}")


if __name__ == "__main__":
    build()
