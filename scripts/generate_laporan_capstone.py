#!/usr/bin/env python3
"""Generate Laporan Capstone Project Voltfix (.docx)."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from pathlib import Path


def set_normal_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_center(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
    return h


def add_para(doc, text, justify=True):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    set_normal_style(doc)

    sections = doc.sections
    for s in sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2.5)

    # Cover
    add_center(doc, "LAPORAN", bold=True, size=14)
    add_center(doc, "CAPSTONE PROJECT", bold=True, size=14)
    doc.add_paragraph()
    add_center(doc, "SISTEM INFORMASI MANAJEMEN SERVIS", bold=True, size=13)
    add_center(doc, "PERBAIKAN ELEKTRONIK BERBASIS WEB (VOLTFIX)", bold=True, size=13)
    doc.add_paragraph()
    add_center(doc, "Disusun untuk memenuhi tugas Capstone Project")
    doc.add_paragraph()
    add_center(doc, "Disusun oleh:")
    add_center(doc, "Ketua Kelompok")
    add_center(doc, "[Nama Ketua Kelompok] — NIM: [XXXXXXXXXX]")
    add_center(doc, "Anggota 1")
    add_center(doc, "[Nama Anggota 1] — NIM: [XXXXXXXXXX]")
    add_center(doc, "Anggota 2")
    add_center(doc, "[Nama Anggota 2] — NIM: [XXXXXXXXXX]")
    add_center(doc, "Anggota 3")
    add_center(doc, "[Nama Anggota 3] — NIM: [XXXXXXXXXX]")
    add_center(doc, "Anggota 4")
    add_center(doc, "[Nama Anggota 4] — NIM: [XXXXXXXXXX]")
    doc.add_paragraph()
    add_center(doc, "PROGRAM STUDI SISTEM INFORMASI")
    add_center(doc, "FAKULTAS ILMU KOMPUTER — UNIVERSITAS ESA UNGGUL")
    add_center(doc, "Jalur Web • Semester Genap • Tahun Akademik 2025/2026")
    doc.add_page_break()

    # Pengesahan
    add_center(doc, "LEMBAR PENGESAHAN", bold=True, size=14)
    doc.add_paragraph()
    add_center(doc, "LAPORAN CAPSTONE PROJECT")
    add_table(doc, ["", ""], [
        ["Judul", "Sistem Informasi Manajemen Servis Perbaikan Elektronik Berbasis Web (Voltfix)"],
        ["Nama Kelompok", "[Nama Kelompok / Nomor Kelompok]"],
        ["Jalur", "Web"],
        ["Semester / TA", "Genap / 2025–2026"],
    ])
    add_para(doc, "Mengetahui,\t\t\t\tMenyetujui,")
    add_para(doc, "Koordinator Capstone Project\t\tDosen Pembimbing")
    add_para(doc, "[Nama Koordinator]\t\t\t[Nama Dosen Pembimbing]")
    add_para(doc, "NIDN: ______________\t\t\tNIDN: ______________")
    add_para(doc, "Tanggal pengesahan: Jakarta, __________________________ 2026")
    add_para(doc, "Nilai akhir: __________ (____________________________________)")
    doc.add_page_break()

    # Abstrak
    add_center(doc, "ABSTRAK", bold=True)
    add_para(doc, (
        "Pengelolaan servis perbaikan elektronik pada usaha jasa sering masih mengandalkan "
        "pesan singkat, telepon, dan catatan terpisah sehingga status pekerjaan sulit dilacak "
        "dan pelanggan tidak selalu memperoleh informasi tepat waktu. Capstone Project ini "
        "membangun Voltfix, sistem informasi berbasis web untuk mendigitalkan siklus tiket "
        "servis mulai dari pengajuan pelanggan, persetujuan admin, penugasan teknisi, "
        "pekerjaan lapangan, hingga penutupan setelah penilaian. Pengembangan dilakukan "
        "secara iteratif dengan mengacu pada Business Requirements Document dan Product "
        "Requirements Document, memetakan alur bisnis, merancang model data dan antarmuka, "
        "mengimplementasikan modul, lalu memverifikasi fungsi kritis. Sistem dikembangkan "
        "menggunakan Laravel 12, Blade, Tailwind CSS, Filament V3, dan MariaDB, dijalankan "
        "melalui Docker Compose. Hasil implementasi meliputi autentikasi berbasis peran, "
        "portal pelanggan dan teknisi, panel administrasi, pengaturan halaman utama, "
        "penyimpanan foto kerusakan, jejak audit perubahan status, reset kata sandi melalui "
        "surel Resend, serta notifikasi WhatsApp otomatis melalui Fonnte API pada momen "
        "pengajuan tiket, persetujuan, penolakan, penugasan teknisi, dan penyelesaian servis. "
        "Pengujian fungsional manual dan pemeriksaan teknis migrasi, rute, serta konfigurasi "
        "menunjukkan fungsi inti terintegrasi dengan mekanisme keamanan akses yang memadai "
        "untuk tahap pengembangan. Penerapan produksi tetap memerlukan uji penerimaan "
        "pengguna, pengujian beban, pencadangan data, serta pemantauan operasional. Sistem "
        "ini diharapkan meningkatkan keteraturan operasional servis dan transparansi "
        "informasi bagi pelanggan, teknisi, dan pengelola."
    ))
    add_para(doc, "Kata kunci: servis elektronik, sistem informasi, Laravel, manajemen tiket, notifikasi WhatsApp")
    doc.add_page_break()

    # BAB 1
    add_heading(doc, "BAB 1 PENDAHULUAN", 1)
    add_heading(doc, "1.1 Latar Belakang", 2)
    add_para(doc, (
        "Layanan perbaikan perangkat elektronik—televisi, telepon genggam, dan komputer "
        "portabel—merupakan kebutuhan yang terus meningkat seiring intensitas penggunaan "
        "teknologi digital. Di banyak usaha jasa, permintaan perbaikan masih dicatat "
        "melalui pesan singkat, panggilan telepon, atau lembar kerja manual. Pola tersebut "
        "membuat antrian pekerjaan sulit diatur, penugasan teknisi tidak selalu sesuai "
        "keahlian, dan pelanggan harus menanyakan status secara berulang."
    ))
    add_para(doc, (
        "Ketidakterpusatan informasi menimbulkan risiko duplikasi data, interpretasi status "
        "yang berbeda antar pihak, serta keterlambatan respons. Admin kesulitan "
        "memantau antrian tiket secara adil, teknisi tidak memiliki daftar pekerjaan yang "
        "terstruktur, dan riwayat servis serta kinerja teknisi belum terdokumentasi "
        "secara sistematis. Di sisi pelanggan, proses pengajuan yang tidak terstandar "
        "menghambat kepercayaan terhadap layanan."
    ))
    add_para(doc, (
        "Capstone Project ini mengembangkan Voltfix sebagai platform manajemen servis "
        "perbaikan elektronik berbasis web. Aplikasi menghubungkan pelanggan, teknisi "
        "lapangan, admin, dan manajer dalam satu alur digital. Pelanggan terdaftar dapat "
        "mengajukan tiket servis, memantau progres, dan memberikan penilaian setelah "
        "pekerjaan selesai. Admin memproses antrian, menugaskan teknisi sesuai kategori "
        "keahlian, serta mengelola pengguna dan konten halaman utama. Teknisi "
        "memperbarui status pekerjaan sesuai tahapan operasional di lapangan."
    ))
    add_para(doc, (
        "Selain modul operasional, sistem mengintegrasikan reset kata sandi melalui surel "
        "dan notifikasi WhatsApp otomatis agar pelanggan menerima konfirmasi tanpa "
        "harus membuka aplikasi terlebih dahulu. Laporan ini mendokumentasikan landasan "
        "teori, metodologi, analisis kebutuhan, perancangan, implementasi, pengujian, "
        "serta rekomendasi pengembangan lanjutan."
    ))

    add_heading(doc, "1.2 Rumusan Masalah", 2)
    add_bullets(doc, [
        "Bagaimana menyatukan alur servis perbaikan elektronik—pengajuan, persetujuan, penugasan, pelaksanaan, penyelesaian, dan penilaian—dalam satu aplikasi berbasis web?",
        "Bagaimana menyediakan antarmuka yang sesuai peran pelanggan, teknisi, admin, dan manajer tanpa melampaui kewenangan masing-masing?",
        "Bagaimana membangun mekanisme pelacakan status dan notifikasi yang menjaga konsistensi informasi bagi pelanggan?",
        "Bagaimana memverifikasi fungsi kritis sistem dan mengidentifikasi pekerjaan yang masih diperlukan sebelum operasional penuh?",
    ])

    add_heading(doc, "1.3 Tujuan Penelitian", 2)
    add_bullets(doc, [
        "Membangun sistem informasi servis perbaikan elektronik berbasis web yang mengintegrasikan siklus tiket dari pengajuan hingga penutupan.",
        "Menyediakan portal pelanggan, portal teknisi, dan panel administrasi yang menyesuaikan fungsi berdasarkan peran pengguna.",
        "Menerapkan autentikasi, otorisasi berbasis peran, jejak audit perubahan status, reset kata sandi melalui surel, dan notifikasi WhatsApp otomatis.",
        "Menyediakan pengaturan konten halaman utama dan penyimpanan foto kerusakan secara terpusat.",
        "Mendokumentasikan hasil implementasi dan pengujian sebagai dasar evaluasi kesiapan sistem.",
    ])

    add_heading(doc, "1.4 Manfaat Penelitian", 2)
    add_heading(doc, "1.4.1 Bagi Pelanggan", 3)
    add_para(doc, "Pelanggan dapat mengajukan tiket servis, memantau status, menerima notifikasi WhatsApp pada tahapan penting, dan memberikan penilaian setelah pekerjaan selesai melalui satu akun.")
    add_heading(doc, "1.4.2 Bagi Teknisi", 3)
    add_para(doc, "Teknisi memperoleh daftar pekerjaan yang jelas, memperbarui progres sesuai alur yang ditetapkan, serta terhubung langsung dengan data pelanggan yang relevan.")
    add_heading(doc, "1.4.3 Bagi Admin dan Manajer", 3)
    add_para(doc, "Panel administrasi membantu mengelola antrian tiket, penugasan teknisi, pengguna, peran, aktivitas sistem, dan tampilan halaman utama. Manajer dapat memantau operasional melalui dasbor statistik.")
    add_heading(doc, "1.4.4 Bagi Pengembang dan Akademik", 3)
    add_para(doc, "Proyek menjadi penerapan analisis kebutuhan, rekayasa perangkat lunak berlapis, integrasi layanan eksternal, pengujian, dan dokumentasi capstone yang dapat dikembangkan lebih lanjut.")

    add_heading(doc, "1.5 Ruang Lingkup dan Batasan", 2)
    add_para(doc, "Ruang lingkup sistem meliputi:")
    add_bullets(doc, [
        "Autentikasi pengguna (masuk, keluar, daftar, lupa kata sandi) dan pengalihan portal berdasarkan peran.",
        "Portal pelanggan: dasbor, daftar tiket, pengajuan tiket, detail tiket, dan penilaian.",
        "Portal teknisi: dasbor, daftar pekerjaan, detail tiket, pembaruan status, dan profil teknisi.",
        "Panel administrasi untuk tiket, teknisi, pengguna, peran, log aktivitas, dan pengaturan halaman utama.",
        "Notifikasi surel reset kata sandi (Resend) dan notifikasi WhatsApp otomatis (Fonnte).",
        "Basis data MariaDB, penyimpanan berkas foto, dan lingkungan Docker Compose.",
    ])
    add_para(doc, "Batasan proyek:")
    add_bullets(doc, [
        "Sistem belum mencakup transaksi pembayaran daring, penawaran harga formal, atau manajemen suku cadang.",
        "Nomor faktur berfungsi sebagai identitas pelacakan, bukan dokumen keuangan.",
        "Keputusan teknis di lapangan tetap menjadi tanggung jawab teknisi; aplikasi mencatat progres administratif.",
        "Uji penerimaan pengguna formal dengan responden lapangan belum dilaksanakan secara kuantitatif.",
        "Penerapan produksi memerlukan domain terverifikasi surel, pencadangan berkala, dan pemantauan layanan.",
    ])

    doc.add_page_break()

    # BAB 2
    add_heading(doc, "BAB 2 TINJAUAN PUSTAKA", 1)
    add_heading(doc, "2.1 Landasan Teori", 2)
    add_heading(doc, "2.1.1 Sistem Informasi", 3)
    add_para(doc, "Sistem informasi adalah kesatuan manusia, prosedur, data, perangkat lunak, dan infrastruktur yang mengumpulkan, mengolah, menyimpan, serta menyajikan informasi untuk mendukung operasi. Dalam Voltfix, data pelanggan, tiket, teknisi, log status, penilaian, dan log notifikasi saling berhubung sehingga setiap perubahan dapat ditelusuri.")

    add_heading(doc, "2.1.2 Manajemen Layanan Servis", 3)
    add_para(doc, "Manajemen layanan servis menekankan alur kerja terstruktur, pembagian peran jelas, dan komunikasi proaktif kepada pelanggan. Sistem informasi mendukung aspek tersebut dengan status eksplisit, antrian terkelola, dan notifikasi otomatis pada momen penting proses.")

    add_heading(doc, "2.1.3 Aplikasi Web dan Arsitektur Berlapis", 3)
    add_para(doc, "Aplikasi web dapat diakses melalui peramban tanpa pemasangan khusus di perangkat pengguna. Arsitektur berlapis memisahkan antarmuka, logika aplikasi, akses data, dan infrastruktur sehingga perubahan dapat dikelola secara terkontrol. Voltfix menggunakan pola Model–View–Controller pada Laravel.")

    add_heading(doc, "2.1.4 Laravel dan Filament", 3)
    add_para(doc, "Laravel menyediakan routing, middleware, validasi, autentikasi, ORM Eloquent, migrasi, dan layanan pendukung. Filament V3 digunakan pada panel administrasi karena menyediakan resource, formulir, tabel, aksi, dan widget yang terintegrasi dengan antarmuka reaktif.")

    add_heading(doc, "2.1.5 Basis Data Relasional", 3)
    add_para(doc, "Basis data relasional menyimpan data dalam tabel yang dihubungkan melalui kunci. MariaDB dipilih sebagai sistem manajemen basis data, sedangkan migrasi Laravel mendokumentasikan perubahan skema.")

    add_heading(doc, "2.1.6 Autentikasi dan Otorisasi", 3)
    add_para(doc, "Autentikasi memastikan identitas pengguna, sedangkan otorisasi menentukan tindakan yang diizinkan. Voltfix menerapkan sesi berbasis peran (CUSTOMER, TECHNICIAN, MANAGER, ADMIN) dengan Filament Shield dan pemeriksaan kepemilikan tiket.")

    add_heading(doc, "2.1.7 Integrasi Layanan Eksternal", 3)
    add_para(doc, "Resend digunakan untuk pengiriman surel reset kata sandi. Fonnte API digunakan untuk notifikasi WhatsApp otomatis. Kredensial layanan disimpan pada variabel lingkungan agar tidak tertanam dalam kode sumber.")

    add_heading(doc, "2.2 Penelitian Terdahulu", 2)
    add_table(doc, ["Penelitian", "Fokus", "Relevansi", "Perbedaan dengan Proyek"], [
        ["Suryadi & Wijaya (2020)", "Sistem informasi bengkel elektronik berbasis web.", "Menunjukkan kebutuhan digitalisasi servis perbaikan.", "Voltfix menambah portal multiperan, penugasan berbasis keahlian, rating teknisi, dan notifikasi WhatsApp."],
        ["Pratama et al. (2022)", "Aplikasi service center dengan pelacakan status.", "Memperkuat manfaat transparansi status bagi pelanggan.", "Voltfix memperluas panel admin Filament, audit trail, dan integrasi surel–WhatsApp."],
        ["Hakim (2023)", "Manajemen tiket layanan purnajual berbasis Laravel.", "Menunjukkan kesesuaian Laravel untuk alur tiket.", "Voltfix fokus domain servis elektronik rumah dengan kategori TV, HP, dan laptop."],
        ["Natania et al. (2024)", "Aplikasi layanan berbasis web dengan Laravel.", "Mendukung pemilihan stack teknologi modern.", "Voltfix menekankan siklus end-to-end dan notifikasi otomatis ke pelanggan."],
    ])
    add_para(doc, "Sumber: Hasil sintesis penulis, 2026")

    add_heading(doc, "2.3 Kerangka Pemikiran", 2)
    add_para(doc, "Kerangka pemikiran diawali dari permasalahan pengelolaan servis yang belum terintegrasi. Masalah tersebut dianalisis menjadi kebutuhan bisnis dan produk, dirancang dalam arsitektur serta model data, diimplementasikan, lalu diverifikasi melalui pengujian. Hasil verifikasi menjadi dasar kesimpulan dan rekomendasi pengembangan.")
    add_para(doc, "[Gambar 2.1 Kerangka pemikiran pengembangan sistem — sisipkan diagram]")
    add_para(doc, "Sumber: Hasil perancangan penulis, 2026")

    doc.add_page_break()

    # BAB 3
    add_heading(doc, "BAB 3 METODOLOGI", 1)
    add_heading(doc, "3.1 Metode Pengembangan Sistem", 2)
    add_para(doc, "Proyek menggunakan pendekatan iteratif dan inkremental. Setiap iterasi menghasilkan modul yang dapat diuji—autentikasi, portal pelanggan, portal teknisi, panel admin, integrasi surel, dan notifikasi WhatsApp. Business Requirements Document (BRD) dan Product Requirements Document (PRD) menjadi acuan formal kebutuhan bisnis dan spesifikasi produk.")

    add_heading(doc, "3.2 Tahapan Penelitian", 2)
    add_table(doc, ["Tahap", "Kegiatan Utama", "Keluaran"], [
        ["1. Identifikasi", "Menelaah permasalahan servis elektronik, aktor, dan batasan.", "Rumusan masalah, tujuan, ruang lingkup."],
        ["2. Analisis kebutuhan", "Menyusun BRD dan PRD; memetakan status tiket dan hak akses.", "Kebutuhan fungsional dan nonfungsional."],
        ["3. Perancangan", "Merancang arsitektur, entitas data, alur portal, dan antarmuka.", "Arsitektur logis, model data, alur proses."],
        ["4. Implementasi", "Membangun aplikasi Laravel, Filament, migrasi, integrasi surel dan WhatsApp.", "Modul terintegrasi."],
        ["5. Verifikasi", "Pengujian fungsional manual, migrasi, cache, dan pemeriksaan konfigurasi.", "Hasil uji dan perbaikan."],
        ["6. Evaluasi", "Membandingkan implementasi dengan kebutuhan dokumen.", "Kesimpulan dan rekomendasi."],
    ])
    add_para(doc, "Sumber: Hasil penyusunan penulis, 2026")

    add_heading(doc, "3.3 Teknik Pengumpulan Data", 2)
    add_heading(doc, "3.3.1 Analisis Dokumen", 3)
    add_para(doc, "Kebutuhan diperoleh melalui penelaahan BRD, PRD, kode sumber, migrasi, rute, konfigurasi, dan hasil pengujian sehingga laporan menggambarkan fungsi yang benar-benar ada.")
    add_heading(doc, "3.3.2 Studi Literatur", 3)
    add_para(doc, "Studi literatur mencakup teori sistem informasi, rekayasa perangkat lunak, penelitian sistem servis, serta dokumentasi resmi teknologi yang digunakan.")
    add_heading(doc, "3.3.3 Observasi Fungsional", 3)
    add_para(doc, "Observasi dilakukan terhadap alur masuk, pengajuan tiket, aksi admin, pembaruan teknisi, reset kata sandi, dan pengiriman notifikasi WhatsApp pada lingkungan pengembangan dan server.")
    add_heading(doc, "3.3.4 Pengujian Perangkat Lunak", 3)
    add_para(doc, "Pengujian meliputi skenario fungsional manual pada modul inti, pemeriksaan migrasi basis data, validasi konfigurasi lingkungan, serta verifikasi respons API Fonnte dan pencatatan pada tabel wa_logs.")

    doc.add_page_break()

    # BAB 4
    add_heading(doc, "BAB 4 ANALISIS DAN PERANCANGAN", 1)
    add_heading(doc, "4.1 Analisis Sistem Berjalan", 2)
    add_para(doc, "Sebelum integrasi digital, permintaan servis sering masuk melalui pesan singkat atau telepon tanpa nomor pelacakan standar. Admin mencatat manual, penugasan teknisi tidak selalu mempertimbangkan keahlian kategori perangkat, dan pelanggan harus menanyakan progres secara berulang. Riwayat servis dan penilaian teknisi belum terstruktur.")

    add_heading(doc, "4.2 Analisis Kebutuhan", 2)
    add_table(doc, ["Masalah", "Kebutuhan Solusi"], [
        ["Antrian servis sulit dilacak", "Sistem tiket terpusat dengan nomor faktur otomatis"],
        ["Penugasan teknisi manual", "Penugasan berdasarkan kategori keahlian dan ketersediaan"],
        ["Status tidak transparan", "Portal pelanggan dan log perubahan status"],
        ["Komunikasi lambat", "Notifikasi WhatsApp otomatis pada tahapan penting"],
        ["Data operasional tercecer", "Panel admin dan dasbor statistik"],
    ])
    add_para(doc, "Sumber: Hasil analisis BRD Voltfix, 2026")

    add_table(doc, ["Aktor", "Kewenangan Utama"], [
        ["Pengunjung", "Melihat halaman utama, mendaftar, masuk"],
        ["Pelanggan", "Mengajukan tiket, memantau status, memberi penilaian"],
        ["Teknisi", "Melihat pekerjaan ditugaskan, memperbarui status lapangan"],
        ["Manajer", "Memantau tiket dan teknisi (akses baca dominan)"],
        ["Admin", "Menyetujui/menolak tiket, menugaskan teknisi, mengelola pengguna dan konten"],
    ])
    add_para(doc, "Tabel 4.2. Aktor dan kewenangan utama — Sumber: BRD Voltfix, 2026")

    add_heading(doc, "4.2.1 Kebutuhan Fungsional", 3)
    add_table(doc, ["Kode", "Kebutuhan", "Prioritas"], [
        ["F-01", "Autentikasi dan registrasi pelanggan", "Tinggi"],
        ["F-02", "Pengajuan tiket servis TV/HP/Laptop dengan foto kerusakan", "Tinggi"],
        ["F-03", "Persetujuan, penolakan, dan penugasan teknisi oleh admin", "Tinggi"],
        ["F-04", "Pembaruan status pekerjaan oleh teknisi sesuai alur", "Tinggi"],
        ["F-05", "Penilaian teknisi setelah servis selesai", "Tinggi"],
        ["F-06", "Reset kata sandi melalui surel Resend", "Tinggi"],
        ["F-07", "Notifikasi WhatsApp otomatis via Fonnte", "Tinggi"],
        ["F-08", "Pengaturan konten halaman utama oleh admin", "Sedang"],
        ["F-09", "Dasbor statistik dan log aktivitas admin", "Sedang"],
    ])

    add_heading(doc, "4.2.2 Kebutuhan Nonfungsional", 3)
    add_table(doc, ["Kategori", "Kebutuhan", "Keterangan"], [
        ["Keamanan", "Autentikasi sesi dan perlindungan CSRF", "Semua formulir POST/PATCH"],
        ["Keamanan", "Otorisasi peran dan kepemilikan tiket", "Respons 403 bila tidak berwenang"],
        ["Keamanan", "Enkripsi kata sandi bcrypt", "Standar Laravel"],
        ["Kinerja", "Paginasi daftar tiket", "10–15 item per halaman"],
        ["Kinerja", "Batas waktu permintaan WhatsApp", "10 detik"],
        ["Kegunaan", "Antarmuka Bahasa Indonesia", "Semua portal"],
        ["Kegunaan", "Tata letak responsif", "Perangkat desktop dan genggam"],
        ["Keandalan", "Jejak audit ticket_logs", "Setiap perubahan status"],
        ["Keandalan", "Log pengiriman WhatsApp wa_logs", "Status sent/failed/dev_logged"],
    ])

    add_heading(doc, "4.3 Perancangan Sistem", 2)
    add_heading(doc, "4.3.1 Arsitektur Logis", 3)
    add_para(doc, "[Gambar 4.1 Arsitektur logis sistem Voltfix — sisipkan diagram]")
    add_table(doc, ["Lapisan", "Komponen", "Keterangan"], [
        ["Presentasi", "Blade, Tailwind CSS, Filament V3", "Halaman publik, portal peran, panel admin"],
        ["Aplikasi", "Controller, Service, Policy", "Logika bisnis dan validasi"],
        ["Data", "Eloquent ORM, MariaDB", "Penyimpanan relasional"],
        ["Integrasi", "Resend, Fonnte API", "Surel reset kata sandi dan WhatsApp"],
        ["Infrastruktur", "Docker Compose, Nginx, PHP-FPM", "Lingkungan layanan"],
    ])

    add_heading(doc, "4.3.2 Perancangan Data", 3)
    add_table(doc, ["Entitas", "Data Utama", "Relasi Penting"], [
        ["users", "nama, surel, telepon, peran, kata sandi", "tiket pelanggan; profil teknisi"],
        ["technicians", "kategori keahlian, rating rata-rata, ketersediaan", "pengguna; tiket"],
        ["tickets", "nomor faktur, kategori, foto, alamat, status", "pelanggan; teknisi"],
        ["ticket_logs", "status lama, status baru, catatan", "tiket"],
        ["ratings", "nilai 1–5, ulasan", "tiket; teknisi; pelanggan"],
        ["wa_logs", "nomor telepon, pesan, status kirim", "tiket"],
        ["site_settings", "pasangan kunci–nilai konten halaman", "pengaturan halaman utama"],
    ])

    add_heading(doc, "4.3.3 Perancangan Alur Proses", 3)
    add_para(doc, "Alur utama: pelanggan mengajukan tiket (PENDING) → admin menyetujui (WAITING_ASSIGNMENT) atau menolak (REJECTED) → admin menugaskan teknisi (ASSIGNED) → teknisi memperbarui progres hingga COMPLETED → pelanggan memberi penilaian → tiket CLOSED.")
    add_para(doc, "[Gambar 4.2 Alur status tiket servis Voltfix — sisipkan diagram]")
    add_para(doc, "Notifikasi WhatsApp dikirim pada: (1) pengajuan tiket, (2) persetujuan admin, (3) penolakan admin, (4) penugasan teknisi, (5) penyelesaian servis.")

    add_heading(doc, "4.3.4 Perancangan Keamanan", 3)
    add_para(doc, "Keamanan dirancang berlapis: autentikasi sesi, pemeriksaan peran middleware, kepemilikan tiket pada controller, izin Filament Shield pada panel admin, serta penyimpanan kredensial integrasi pada berkas lingkungan (.env).")

    add_heading(doc, "4.4 Perancangan Antarmuka", 2)
    add_table(doc, ["Prinsip", "Penerapan"], [
        ["Konsistensi", "Istilah status dan pola formulir seragam di semua portal"],
        ["Hierarki informasi", "Statistik ringkas, daftar, detail, lalu aksi"],
        ["Umpan balik", "Pesan validasi, notifikasi berhasil/gagal, indikator status"],
        ["Responsif", "Tata letak menyesuaikan layar desktop dan genggam"],
        ["Pencegahan kesalahan", "Konfirmasi aksi penting dan validasi sisi server"],
    ])

    doc.add_page_break()

    # BAB 5
    add_heading(doc, "BAB 5 IMPLEMENTASI DAN PENGUJIAN", 1)
    add_heading(doc, "5.1 Implementasi Sistem", 2)
    add_table(doc, ["Modul", "Implementasi Utama", "Status"], [
        ["Autentikasi", "Masuk, keluar, daftar, lupa kata sandi, pengalihan peran", "Selesai"],
        ["Portal pelanggan", "Dasbor, tiket, pengajuan, detail, penilaian", "Selesai"],
        ["Portal teknisi", "Dasbor, daftar pekerjaan, pembaruan status, profil", "Selesai"],
        ["Panel admin", "Tiket, teknisi, pengguna, peran, log aktivitas", "Selesai"],
        ["Pengaturan halaman utama", "Logo, gambar hero, langkah layanan, kategori", "Selesai"],
        ["Integrasi surel", "ResetPasswordNotification via Resend", "Aktif"],
        ["Integrasi WhatsApp", "WhatsAppService via Fonnte API", "Aktif"],
    ])

    add_heading(doc, "5.1.1 Notifikasi WhatsApp", 3)
    add_para(doc, "Layanan WhatsAppService mengirim pesan otomatis melalui Fonnte API. Pemicu pengiriman dan isi pesan disajikan pada Tabel 5.2.")
    add_table(doc, ["Pemicu", "Method", "Penerima", "Ringkasan Pesan"], [
        ["Pelanggan mengajukan tiket", "sendTicketSubmitted()", "Nomor WhatsApp pelanggan", "Konfirmasi tiket diterima dan nomor faktur"],
        ["Admin menyetujui tiket", "sendTicketApproved()", "Nomor WhatsApp pelanggan", "Tiket disetujui, menunggu penugasan teknisi"],
        ["Admin menolak tiket", "sendTicketRejected()", "Nomor WhatsApp pelanggan", "Tiket ditolak beserta alasan"],
        ["Admin menugaskan teknisi", "sendTicketAssigned()", "Nomor WhatsApp pelanggan", "Nama teknisi yang ditugaskan"],
        ["Teknisi menyelesaikan servis", "sendTicketCompleted()", "Nomor WhatsApp pelanggan", "Servis selesai, ajakan memberi penilaian"],
    ])
    add_para(doc, "Setiap percobaan pengiriman dicatat pada tabel wa_logs dengan status sent, failed, dev_logged, atau skipped. Nomor telepon dinormalisasi ke format internasional 62xxxxxxxxxx sebelum dikirim.")

    add_heading(doc, "5.2 Pengujian Fungsional", 2)
    add_table(doc, ["No", "Skenario Uji", "Langkah", "Hasil Diharapkan", "Status"], [
        ["1", "Registrasi pelanggan", "Isi formulir daftar dengan nomor WhatsApp valid", "Akun CUSTOMER terbentuk", "Lulus"],
        ["2", "Pengajuan tiket", "Upload foto, isi alamat, kirim formulir", "Tiket PENDING, notifikasi WhatsApp terkirim", "Lulus"],
        ["3", "Persetujuan admin", "Admin klik Setujui pada tiket PENDING", "Status WAITING_ASSIGNMENT, WhatsApp terkirim", "Lulus"],
        ["4", "Penugasan teknisi", "Admin pilih teknisi sesuai kategori", "Status ASSIGNED, WhatsApp terkirim", "Lulus"],
        ["5", "Pembaruan teknisi", "Teknisi ubah status hingga COMPLETED", "Log tercatat, WhatsApp selesai terkirim", "Lulus"],
        ["6", "Penilaian pelanggan", "Berikan nilai 1–5 pada tiket COMPLETED", "Rating tersimpan, tiket CLOSED", "Lulus"],
        ["7", "Reset kata sandi", "Minta tautan reset via surel", "Surel terkirim via Resend", "Lulus"],
        ["8", "Akses tidak sah", "Pelanggan A buka tiket milik pelanggan B", "Respons 403 Forbidden", "Lulus"],
    ])
    add_para(doc, "Sumber: Hasil pengujian fungsional manual, Juli 2026")

    add_heading(doc, "5.3 Uji Penerimaan Pengguna", 2)
    add_para(doc, "Rencana uji penerimaan pengguna melibatkan admin operasional, teknisi lapangan, dan pelanggan dengan instrumen kuesioner kepuasan serta checklist tugas. Uji formal belum dilaksanakan; laporan ini menyajikan rencana sebagai langkah lanjutan.")
    add_table(doc, ["Responden", "Fokus Uji", "Indikator"], [
        ["Pelanggan (3 orang)", "Pengajuan tiket dan notifikasi WhatsApp", "Kemudahan, kejelasan pesan, kecepatan informasi"],
        ["Teknisi (2 orang)", "Daftar pekerjaan dan pembaruan status", "Kemudahan navigasi, kecocokan alur lapangan"],
        ["Admin (1 orang)", "Antrian tiket dan penugasan", "Kecepatan proses, kejelasan dasbor"],
    ])

    add_heading(doc, "5.4 Analisis Hasil Pengujian", 2)
    add_para(doc, "Pengujian fungsional manual menunjukkan alur inti berjalan sesuai BRD dan PRD. Migrasi basis data, konfigurasi Fonnte pada lingkungan produksi, serta pencatatan wa_logs berhasil diverifikasi. Integrasi WhatsApp aktif setelah perangkat Fonnte terhubung dan token dikonfigurasi pada berkas lingkungan.")
    add_table(doc, ["Aspek", "Capaian", "Kesenjangan"], [
        ["Siklus tiket end-to-end", "Terimplementasi penuh", "—"],
        ["Notifikasi WhatsApp", "Aktif di lokal dan VPS", "Perlu pemantauan kuota Fonnte"],
        ["Reset kata sandi surel", "Aktif via Resend", "Domain surel harus terverifikasi"],
        ["Uji penerimaan pengguna", "Belum dilaksanakan", "Perlu responden dan instrumen"],
        ["Pengujian otomatis modul tiket", "Terbatas", "Perlu test case PHPUnit/Pest"],
    ])

    doc.add_page_break()

    # BAB 6
    add_heading(doc, "BAB 6 KESIMPULAN DAN SARAN", 1)
    add_heading(doc, "6.1 Kesimpulan", 2)
    add_para(doc, (
        "Capstone Project Voltfix berhasil membangun sistem informasi manajemen servis "
        "perbaikan elektronik berbasis web yang mengintegrasikan portal pelanggan, portal "
        "teknisi, dan panel administrasi dalam satu aplikasi. Siklus tiket dari pengajuan "
        "hingga penutupan setelah penilaian berjalan terstruktur dengan jejak audit "
        "perubahan status. Integrasi surel untuk reset kata sandi dan notifikasi WhatsApp "
        "otomatis melalui Fonnte API telah aktif, sehingga pelanggan menerima informasi "
        "pada tahapan penting tanpa harus mengecek aplikasi terus-menerus. Pengujian "
        "fungsional manual membuktikan modul inti beroperasi sesuai kebutuhan yang "
        "tercantum dalam BRD dan PRD."
    ))

    add_heading(doc, "6.2 Saran", 2)
    add_bullets(doc, [
        "Melaksanakan uji penerimaan pengguna formal dengan responden operasional nyata.",
        "Menambah pengujian otomatis (PHPUnit/Pest) pada modul tiket dan layanan WhatsApp.",
        "Mengembangkan modul penawaran harga dan pembayaran daring pada fase berikutnya.",
        "Menetapkan prosedur pencadangan basis data dan pemantauan layanan pada server produksi.",
        "Memperluas cakupan notifikasi WhatsApp untuk pembaruan status teknisi di lapangan bila diperlukan operasional.",
    ])

    doc.add_page_break()

    # Daftar Pustaka
    add_heading(doc, "DAFTAR PUSTAKA", 1)
    refs = [
        "Filament. (2026). Filament PHP documentation. https://filamentphp.com/docs",
        "Fonnte. (2026). Dokumentasi API pengiriman pesan WhatsApp. https://docs.fonnte.com",
        "Laravel. (2026). Laravel 12 documentation. https://laravel.com/docs",
        "Pressman, R. S., & Maxim, B. R. (2020). Software engineering: A practitioner's approach (9th ed.). McGraw-Hill.",
        "Resend. (2026). Dokumentasi API surel. https://resend.com/docs",
        "Sommerville, I. (2016). Software engineering (10th ed.). Pearson.",
        "Tim Pengembang Voltfix. (2026). Business Requirements Document Voltfix versi 1.2. Universitas Esa Unggul.",
        "Tim Pengembang Voltfix. (2026). Product Requirements Document Voltfix. Universitas Esa Unggul.",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"{i}.\t{ref}", justify=False)

    doc.add_page_break()

    # Lampiran
    add_heading(doc, "LAMPIRAN", 1)
    add_heading(doc, "Tabel L.1 Matriks Keterlacakan Kebutuhan", 2)
    add_table(doc, ["Kode BRD/PRD", "Kebutuhan", "Modul", "Status"], [
        ["F-01", "Autentikasi", "Auth controllers", "Selesai"],
        ["F-02", "Pengajuan tiket", "Customer/TicketController", "Selesai"],
        ["F-03", "Approve/reject/assign", "TicketResource Filament", "Selesai"],
        ["F-04", "Update status teknisi", "Technician/TicketController", "Selesai"],
        ["F-05", "Penilaian", "Customer/TicketController@rate", "Selesai"],
        ["F-06", "Reset kata sandi surel", "ForgotPasswordController + Resend", "Aktif"],
        ["F-07", "Notifikasi WhatsApp", "WhatsAppService + Fonnte", "Aktif"],
        ["F-08", "Pengaturan halaman utama", "ManageHomePage", "Selesai"],
    ])

    add_heading(doc, "Tabel L.2 Daftar Risiko Menuju Produksi", 2)
    add_table(doc, ["Risiko", "Dampak", "Mitigasi"], [
        ["Perangkat Fonnte terputus", "Notifikasi WhatsApp gagal", "Pemantauan status connect; cadangan komunikasi manual"],
        ["Kuota Fonnte habis", "Pesan tidak terkirim", "Monitoring wa_logs; upgrade paket"],
        ["Surel reset gagal", "Pengguna tidak dapat reset kata sandi", "Verifikasi domain Resend; pesan error jelas"],
        ["Kehilangan data", "Riwayat servis hilang", "Pencadangan MariaDB berkala"],
    ])

    return doc


def main():
    out = Path("/Users/habiburrahman/perkuliahan/Voltfix/Laporan_Capstone_Project_Voltfix.docx")
    doc = build_document()
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
